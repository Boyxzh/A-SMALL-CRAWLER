import requests
import re
from datetime import datetime
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import logging

# 配置日志记录
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.FileHandler('crawler.log'), logging.StreamHandler()]
)

HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
}
BASE_URL = 'https://news.cri.cn'

def get_article_urls(page_type, page=1):
    """获取文章链接"""
    try:
        if page_type == "first":
            url = f'{BASE_URL}/guojiruiping'
        else:
            url = f'{BASE_URL}/inc/08b152bb-56a6-4ba4-9b60-8761fa1a1568-{page}.inc'

        response = requests.get(url, headers=HEADERS, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        links = soup.select('div.sftx-list.more-list a[href]')
        return list({link['href'] for link in links if link['href'].strip()})
        
    except Exception as e:
        logging.error(f"获取链接失败: {str(e)}")
        return []

def process_article(url):
    """处理单个文章页面"""
    try:
        # 处理相对路径
        if not url.startswith(('http://', 'https://')):
            article_url = f'{BASE_URL}{url}'
        else:
            article_url = url

        response = requests.get(article_url, headers=HEADERS, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')

        # 提取标题
        title_tag = soup.find('div', class_="list-title")
        if not title_tag:
            raise ValueError("未找到标题元素")
        title = title_tag.get_text(strip=True)


        # 提取时间并格式化
        time = soup.find('div', class_="list-brief")
        time_article_ = time.get_text()
        time_article = time_article_[:10]
         


        # 提取正文内容
        content_tag = soup.find('div', class_="list-abody abody")
        if not content_tag:
            raise ValueError("未找到正文内容")
        content = content_tag.get_text(strip=True)

        # 创建Word文档
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        doc.add_heading(title, level=1)
        doc.add_paragraph(content)

        # 生成安全文件名
        clean_title = re.sub(r'[\\/:*?"<>|]', '', title)[:50]  # 限制标题长度
        filename = f"{time_article}_{clean_title}"
        doc.save(f"{filename}.docx")
        logging.info(f"成功保存：{filename}")

    except Exception as e:
        logging.error(f"处理文章失败：{url}，错误：{str(e)}")

def main():
    while True:
        try:
            total = int(input("请输入需要获取的文章数量："))
            if total <= 0:
                print("请输入正整数")
                continue
                
            collected_urls = []
            page = 1
            
            # 动态获取直到满足数量需求（修复分页逻辑）
            while len(collected_urls) < total:
                # 获取当前页链接
                current_urls = get_article_urls(
                    "first" if page == 1 else "subsequent", 
                    page
                )
                
                if not current_urls:
                    logging.warning("没有更多文章可获取")
                    break
                
                # 去重并添加新链接
                new_urls = [u for u in current_urls if u not in collected_urls]
                collected_urls.extend(new_urls)
                
                logging.info(f"已获取第{page}页，当前总数：{len(collected_urls)}")
                page += 1

            # 截取所需数量
            for url in collected_urls[:total]:
                process_article(url)

            print(f"任务完成！共处理{min(len(collected_urls), total)}篇文章")
            break

        except ValueError:
            print("请输入有效的数字")
        except Exception as e:
            logging.error(f"程序运行异常：{str(e)}")
            break

if __name__ == '__main__':
    main()
